import json
import re
from PyPDF2 import PdfReader
from docx import Document
import requests
import io
from fastapi import HTTPException

def extract_json_objects(text: str):
    """
    Extracts all JSON objects from the given text and returns them as a list of dicts.
    Handles both raw JSON and JSON wrapped in markdown code blocks.
    Also handles multiline strings that need proper escaping.
    """
    import json
    import re
    
    # First, try to extract JSON from markdown code blocks
    code_block_pattern = r'```(?:json)?\s*(\{.*?\})\s*```'
    code_block_matches = re.findall(code_block_pattern, text, re.DOTALL)
    
    json_objects = []
    
    # Try to parse JSON from code blocks first
    for match in code_block_matches:
        try:
            # Try to fix common JSON issues with multiline strings
            fixed_match = fix_multiline_json_strings(match.strip())
            obj = json.loads(fixed_match)
            json_objects.append(obj)
        except json.JSONDecodeError:
            continue
    
    # If no valid JSON found in code blocks, try to find raw JSON objects
    if not json_objects:
        # Simple pattern to match JSON objects - find balanced braces
        def find_json_objects(text):
            results = []
            i = 0
            while i < len(text):
                if text[i] == '{':
                    # Found start of potential JSON object
                    brace_count = 1
                    start = i
                    i += 1
                    
                    while i < len(text) and brace_count > 0:
                        if text[i] == '{':
                            brace_count += 1
                        elif text[i] == '}':
                            brace_count -= 1
                        i += 1
                    
                    if brace_count == 0:
                        # Found complete JSON object
                        json_candidate = text[start:i]
                        results.append(json_candidate)
                else:
                    i += 1
            return results
        
        potential_json_objects = find_json_objects(text)
        
        for json_candidate in potential_json_objects:
            try:
                # Try to fix multiline string issues
                fixed_match = fix_multiline_json_strings(json_candidate)
                obj = json.loads(fixed_match)
                json_objects.append(obj)
            except json.JSONDecodeError:
                continue
    
    return json_objects


def fix_multiline_json_strings(json_string: str) -> str:
    """
    Fixes common issues with multiline strings in JSON responses from AI.
    """
    import re
    
    # Pattern to match string values that might contain unescaped newlines
    def replace_unescaped_newlines(match):
        full_match = match.group(0)
        key = match.group(1)
        value = match.group(2)
        
        # Replace actual newlines with \n
        escaped_value = value.replace('\n', '\\n').replace('\r', '\\r')
        # Fix any double escaping
        escaped_value = escaped_value.replace('\\\\n', '\\n')
        
        return f'"{key}": "{escaped_value}"'
    
    # Pattern to match "key": "value with possible newlines"
    pattern = r'"([^"]+)":\s*"((?:[^"\\]|\\.)*)(?<!\\)"'
    
    # Replace unescaped newlines in string values
    fixed_json = re.sub(pattern, replace_unescaped_newlines, json_string, flags=re.DOTALL)
    
    return fixed_json


def extract_text(file):
    """
    Extract text content from uploaded files (PDF, DOCX, TXT)
    """
    try:
        content = file.file.read()
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            # Handle PDF
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            return text.strip()
            
        elif filename.endswith('.docx'):
            # Handle DOCX
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            text = ''
            for para in doc.paragraphs:
                text += para.text + '\n'
            return text.strip()
            
        elif filename.endswith('.txt'):
            # Handle plain text
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = content.decode('latin-1')
                except UnicodeDecodeError:
                    text = content.decode('utf-8', errors='ignore')
            return text.strip()
            
        else:
            # Fallback: try to read as text
            try:
                text = content.decode('utf-8')
                return text.strip()
            except UnicodeDecodeError:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Unsupported file type or encoding: {filename}"
                )
            
    except Exception as e:
        raise HTTPException(
            status_code=400, 
            detail=f"Failed to extract text from file {file.filename}: {str(e)}"
        )
    finally:
        # Reset file pointer for potential future reads
        file.file.seek(0)


def extract_text_from_url(url: str):
    """Extract text from a file URL (Cloudinary, etc.)"""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        # Get content type
        content_type = response.headers.get('content-type', '').lower()
        
        if 'pdf' in content_type or url.lower().endswith('.pdf'):
            # Handle PDF
            pdf_file = io.BytesIO(response.content)
            reader = PdfReader(pdf_file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() or ''
            return text
            
        elif 'document' in content_type or url.lower().endswith('.docx'):
            # Handle DOCX
            docx_file = io.BytesIO(response.content)
            doc = Document(docx_file)
            text = ''
            for para in doc.paragraphs:
                text += para.text + '\n'
            return text
            
        elif 'text' in content_type or url.lower().endswith('.txt'):
            # Handle plain text
            return response.text
            
        else:
            # Try to read as text anyway
            return response.text
            
    except requests.exceptions.RequestException as e:
        raise HTTPException(
            status_code=400, 
            detail=f"Failed to download file from URL: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=400, 
            detail=f"Failed to extract text from file: {str(e)}"
        )


